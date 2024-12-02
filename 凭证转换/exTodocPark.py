import os
import sys
import win32com.client as win32
import pandas as pd
from docx import Document

# 加载Excel文件
excel_file = r'd:/凭证转换/input/原始文件.xlsx'  # 替换为你的实际Excel文件路径
df = pd.read_excel(excel_file)


def resource_path(relative_path):
    """获取打包后的文件路径"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)


# 按“园区”分组
grouped = df.groupby('园区')

for park, group in grouped:
    # 1. 提取唯一的“劳动合同主体”
    # contract_subjects = group['劳动合同主体'].unique()
    contract_subjects = ""
    default_subject = "上海千服企业管理有限公司；"
    contract_subjects_str = default_subject + '；'.join(map(str, contract_subjects))  # 将所有元素转换为字符串再拼接

    # 2. 提取员工详情并复制为副本，避免SettingWithCopyWarning
    employee_data = group[['员工姓名', '证件号码']].copy()  # 这里添加copy()
    employee_data['序号'] = range(1, len(employee_data) + 1)  # 自动生成序号
    employee_data['证件'] = '身份证'  # 添加“证件”列，值全为“身份证”
    employee_data['职业类别'] = '四类'  # 添加“职业类别”列，值全为“员工”

    # 3. 创建Word文档
    template_path = resource_path('templatepark.docx')
    doc = Document(template_path)

    # 在Word文档的第三行插入劳动合同主体字符串（假设是第3行第2列）
    # cell = doc.tables[0].cell(2, 1)
    # run = cell.paragraphs[0].add_run(contract_subjects_str)
    # run.bold = True  # 设置劳动合同主体为粗体

    # 在文档最后添加员工表格
    doc.add_paragraph("员工信息").runs[0].bold = True  # 添加粗体标题
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('序号').bold = True
    hdr_cells[1].paragraphs[0].add_run('姓名').bold = True
    hdr_cells[2].paragraphs[0].add_run('证件').bold = True
    hdr_cells[3].paragraphs[0].add_run('身份证号码').bold = True
    hdr_cells[4].paragraphs[0].add_run('园区名称').bold = True

    # 填充表格数据
    for index, row in employee_data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['序号'])
        row_cells[1].text = row['员工姓名']
        row_cells[2].text = row['证件']
        row_cells[3].text = row['证件号码']
        # row_cells[4].text = row['职业类别']
        row_cells[4].text = f'{park}'
    # 保存文档，名称为"园区+在保证明.docx"
    default_park = "-在保凭证"
    file_name = f'{park}' + default_park
    file_name = file_name.replace(' ', '_')  # 去除空格，避免文件名不合法
    word_output_path = os.path.join('D:/凭证转换/output/', f'{file_name}.docx')
    doc.save(word_output_path)
    # 将 Word 文件转换为 PDF
    pdf_output_path = word_output_path.replace('.docx', '.pdf')

    # 使用 Word 应用程序将 DOCX 转换为 PDF
    word = win32.Dispatch('Word.Application')
    doc = word.Documents.Open(word_output_path)
    doc.SaveAs(pdf_output_path, FileFormat=17)  # 17 对应于 PDF 格式
    doc.Close()
    word.Quit()

    os.remove(word_output_path)

    print(f'PDF 文件已保存到 {pdf_output_path}')
