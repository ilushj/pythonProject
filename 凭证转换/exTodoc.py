import os
import glob
import sys
import pandas as pd
import re
from docx import Document
from datetime import datetime
from docx.shared import Pt, Inches
from openpyxl.reader.excel import load_workbook
import win32com.client as win32

# 设置要遍历的目录路径
directory_path = r'D:/凭证转换/input/'
# 获取指定目录中的所有 Excel 文件
excel_files = glob.glob(os.path.join(directory_path, '*.xlsx'))


def resource_path(relative_path):
    """获取打包后的文件路径"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)


# 函数：处理每个 Excel 文件并生成 PDF
def process_excel_to_pdf(excel_file_path):
    # 读取 Excel 文件
    wb = load_workbook(excel_file_path, data_only=True)  # data_only读取公式的计算结果
    sheet = wb.active  # 获取当前活动的表

    # 从 Excel 文件中提取变量
    insured_person = sheet['J2'].value  # 被保险人
    policy_number = sheet['D3'].value  # 保单号
    insurance_type = sheet['J3'].value  # 险种类型
    start_date = sheet['D4'].value  # 起保日期
    end_date = sheet['J4'].value  # 终保日期



    # 提取保额
    match = re.search(r'保额(.+?)(保费|$)', sheet['D6'].value)

    if match:
        quota = match.group(1).strip()
    else:
        quota = None  # 或其他默认值，视情况而定
        # 加载 Word 模板
    # 使用 resource_path 获取打包后的模板文件路径
    template_path = resource_path('template.docx')
    doc = Document(template_path)

    # print(f"提取的变量: 被保险人: {insured_person}, 保单号: {policy_number}, 险种类型: {insurance_type}")
    # 添加从 Excel 提取的变量信息
    def add_variable_paragraph(text, font_name='宋体', font_size=10):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)  # 五号字
        run.bold = True  # 设置粗体

    # 添加变量信息
    if insurance_type == '雇主责任险':
        add_variable_paragraph(f'被保险人：{insured_person}')
    elif insurance_type == '团体意外险':
        add_variable_paragraph(f'投保人  ：{insured_person}')
    add_variable_paragraph(f'保单号码：{policy_number}')
    add_variable_paragraph(f'险种类型：{insurance_type}')
    add_variable_paragraph(f'起保日期：{start_date}   终保日期：{end_date}')

    doc.add_paragraph()  # 添加第二个空行

    # 添加表格头部（模仿 Word 示例的表格）
    table = doc.add_table(rows=1, cols=7)

    # 设置表格样式为 "Table Grid"（添加边框线）
    table.style = 'Table Grid'

    # 手动设置每一列的宽度
    table.columns[0].width = Inches(0.3)  # 序号列
    table.columns[1].width = Inches(1.2)  # 姓名列
    table.columns[2].width = Inches(1.6)  # 身份证号码列
    table.columns[3].width = Inches(1.6)  # 变更类型列
    table.columns[4].width = Inches(1.6)  # 生效日期列
    table.columns[5].width = Inches(1.6)  # 到期日期列
    table.columns[6].width = Inches(0.3)  # 保额列

    # 读取 Excel 数据表
    df = pd.read_excel(excel_file_path, dtype={'证件号码': str}, skiprows=6)  # 根据实际情况调整 skiprows

    # 清理列名空格，并确保列名一致
    df.columns = df.columns.str.strip()  # 去掉列名的空格

    # 筛选 "批增/批减" 列中为 "批增" 的数据
    df_filtered = df[df['批增/批减'] == '批增'].copy()

    # 添加表格头部
    header_cells = table.rows[0].cells
    header_cells[0].text = '序号'
    header_cells[1].text = '姓名'
    header_cells[2].text = '身份证号码'
    header_cells[3].text = '变更类型'
    header_cells[4].text = '生效日期'
    header_cells[5].text = '到期日期'
    header_cells[6].text = '保额'

    # 填充表格内容，并生成自增序号
    for i, (index, row) in enumerate(df_filtered.iterrows(), start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)  # 自增序号

        values = [
            row['姓名'].replace(' ', ''),
            row['证件号码'],
            row['批增/批减'],
            row['生效日期'].strftime('%Y-%m-%d') if isinstance(row['生效日期'], datetime) else str(row['生效日期']),
            row['到期日期'].strftime('%Y-%m-%d') if isinstance(row['到期日期'], datetime) else str(row['到期日期']),
            quota
        ]

        # 遍历值，并为每个单元格中的文本设置字体
        for j, value in enumerate(values):
            cells[j + 1].text = str(value)

    # 在添加内容之前插入两个空行
    doc.add_paragraph()  # 添加第一个空行
    doc.add_paragraph()  # 添加第二个空行

    add_variable_paragraph(
        f'*兹经被保险人申请，本公司对上述批改内容予以确认。',
        font_name='宋体', font_size=7
    )
    add_variable_paragraph(
        f'*该保险凭证的盖章件，仅限于员工入场时使用，不用于其他用途，不作为理赔依据，否则本公司有权追究相关的法律责任。',
        font_name='宋体', font_size=7
    )

    # 生成文件名，包含被保险人名称和起保日期
    file_name = f'{insured_person}'.replace(' ', '_')  # 去除空格，避免文件名不合法
    word_output_path = os.path.join('D:/凭证转换/output/', f'{file_name}.docx')

    doc.save(word_output_path)

    # 将 Word 文件转换为 PDF
    pdf_output_path = word_output_path.replace('.docx', '.pdf')

    # 使用 Word 应用程序将 DOCX 转换为 PDF
    word = win32.Dispatch('Word.Application')
    doc = word.Documents.Open(word_output_path)
    doc.SaveAs(pdf_output_path, FileFormat=17)  # 17 对应于 PDF 格式
    doc.Close()
    word.Quit()

    os.remove(word_output_path)

    print(f'PDF 文件已保存到 {pdf_output_path}')


for excel_file in excel_files:
    print(f'正在处理文件：{excel_file}')
    process_excel_to_pdf(excel_file)
